

import hashlib
import json
import logging
import matplotlib.pyplot as plt
import nltk
import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import re
import streamlit as st
import string
import unittest
import spacy

from datetime import datetime
from nltk.corpus import stopwords
from nltk.sentiment import SentimentIntensityAnalyzer
from nltk.stem import WordNetLemmatizer
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import word_tokenize
from sklearn.feature_extraction.text import TfidfVectorizer
from textblob import TextBlob
from wordcloud import WordCloud
from sklearn.cluster import KMeans
import re
import string
from datetime import datetime
from sklearn.feature_extraction.text import TfidfVectorizer
from textblob import TextBlob


nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')
nltk.download('vader_lexicon')
nltk.download('punkt', download_dir='/content')
nltk.download('stopwords', download_dir='/content')
nltk.download('wordnet', download_dir='/content')
nltk.download('vader_lexicon', download_dir='/content')
nltk.data.path.append('/content')

@st.cache_resource
def download_punkt():
    nltk.download('punkt')



# Download NLTK resources at startup
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', force=True)

try:
    nltk.download('vader_lexicon', quiet=True)
    nltk.download('stopwords', quiet=True)
    nltk.download('wordnet', quiet=True)
except Exception as e:
    st.error(f"Error downloading NLTK resources: {str(e)}")

# Define stopwords and lemmatizer
#stop_words = set(stopwords.words('english'))
stop_words = stopwords.words('english')
stop_words += ['…', 'nuclearenergy', '’', 'yes', '核エネルギーの潜在的な危険性は、いくら強調してもし過ぎるということはない。']
lemmatizer = WordNetLemmatizer()

# Security configurations
class SecurityConfig:
    def __init__(self):
        self.ALLOWED_EXTENSIONS = {'csv', 'txt'}
        self.MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB

    def validate_file(self, file):
        if not file:
            return False
        extension = file.name.split('.')[-1].lower()
        return extension in self.ALLOWED_EXTENSIONS and len(file.getvalue()) <= self.MAX_FILE_SIZE

# text processing and analysis
class TextAnalyzer:
    def __init__(self):
        try:
            self.sia = SentimentIntensityAnalyzer()
        except Exception as e:
            st.error(f"Error initializing VADER analyzer: {str(e)}")
            self.sia = None

    def analyze_sentiment(self, text):
        results = {}

        # VADER sentiment
        if self.sia:
            try:
                results['vader'] = self.sia.polarity_scores(str(text))['compound']
            except Exception as e:
                results['vader'] = 0
                logging.error(f"VADER analysis error: {str(e)}")

        # TextBlob sentiment
        try:
            results['textblob'] = TextBlob(str(text)).sentiment.polarity
        except Exception as e:
            results['textblob'] = 0
            logging.error(f"TextBlob analysis error: {str(e)}")

        return results

    def get_basic_stats(self, text):
        """Get basic text statistics"""
        text = str(text)
        return {
            'char_count': len(text),
            'word_count': len(text.split()),
            'sentence_count': len(text.split('.')),
        }

# Data processing and saving utilities
class DataManager:
    @staticmethod
    def save_results(data, filename):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"analysis_results_{timestamp}.json"
        with open(filename, 'w') as f:
            json.dump(data, f)
        return filename




def main():
    st.set_page_config(page_title="Text Analysis Pipeline", layout="wide")
    st.title("Sentiment Data Analysis Pipeline")
    st.sidebar.title("Data Pipeline Steps")

    # Initialize components
    security = SecurityConfig()
    analyzer = TextAnalyzer()

    # Setup logging
    logging.basicConfig(filename='app.log', level=logging.INFO)


    # Simple authentication (accept any password)
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False

        # Create a placeholder for authentication widgets
    auth_placeholder = st.empty()

    if not st.session_state.authenticated:
        with auth_placeholder:
            # Show the password input field
            password = st.text_input("Enter password", type="password")
            if password:
                # Authenticate and clear the widgets
                st.session_state.authenticated = True
                st.success("Authenticated")
                auth_placeholder.empty()

    # Main navigation triggered immediately after authentication
    if st.session_state.authenticated:
        pages = {
            "Data Loading": render_data_loading,
            "Data Display": render_data_display,
            "Data Cleaning": render_data_cleaning,
            "Data Preprocessing": render_data_preprocessing,
            "Data Analysis Methods and Visualization": render_analysis,
            "Generate Reports": render_reports,
            "Help": render_help
        }

        page = st.sidebar.radio("Navigation", list(pages.keys()))
        pages[page](st.session_state, analyzer, security)



def render_data_loading(session_state, analyzer, security):
    st.title("Data Loading")  # Main heading for the page
    st.subheader("Choose data source")

    # Add options for data input
    data_source = st.radio("Choose data source", ["Upload File", "Text Input", "URL Input"])

    # Option 1: Upload File
    if data_source == "Upload File":
        uploaded_file = st.file_uploader("Upload CSV/TXT file", type=["csv", "txt"])
        if uploaded_file and security.validate_file(uploaded_file):
            try:
                data = pd.read_csv(uploaded_file)
                session_state.data = data
                st.success("File uploaded successfully!")
                st.write(data.head())  # Display a preview of the data
            except Exception as e:
                st.error(f"Error processing the file: {str(e)}")
        elif uploaded_file:
            st.error("Invalid file. Please upload a valid CSV or TXT file.")

    # Option 2: Text Input
    elif data_source == "Text Input":
        text_input = st.text_area("Enter text data for analysis")
        if text_input:
            try:
                session_state.data = pd.DataFrame([{"text": text_input}])
                st.success("Text input received!")
                st.write(session_state.data)  # Display a preview of the data
            except Exception as e:
                st.error(f"Error processing text input: {str(e)}")

    # Option 3: URL Input
    elif data_source == "URL Input":
        url = st.text_input("Enter URL")
        if url and st.button("Fetch Data"):
            try:
                data = pd.read_csv(url)
                session_state.data = data
                st.success("Data fetched successfully from the URL!")
                st.write(data.head())  # Display a preview of the data
            except Exception as e:
                st.error(f"Failed to fetch data from URL: {str(e)}")

def render_data_display(session_state, analyzer, security):
    st.title("Data Display")

    # Check if data is loaded
    if 'data' in session_state and session_state.data is not None:
        # Radio buttons for display options
        display_option = st.radio(
            "Choose display format",
            ["Raw Data", "Summary Statistics", "Column Info","Summary Of Missing Values","Summary Of Duplicates"]
        )

        # Display raw data
        if display_option == "Raw Data":
            st.dataframe(session_state.data)

        # Display summary statistics
        elif display_option == "Summary Statistics":
            st.write(session_state.data.describe())

        # Display column information
        elif display_option == "Column Info":
            from io import StringIO
            buffer = StringIO()
            session_state.data.info(buf=buffer)
            info_output = buffer.getvalue()
            st.text(info_output)

        # Display count of missing values in each column
        elif display_option == "Summary Of Missing Values":
            missing_values = session_state.data.isna().sum()
            st.write(missing_values)

        # Display count of duplicate values
        elif display_option == "Summary Of Duplicates":
            duplicate_values = session_state.data.duplicated().sum()
            st.write(f"Number of duplicate rows: {duplicate_values}")

    else:
        # Warning if no data is loaded
        st.warning("Please load data first.")

def render_data_cleaning(session_state, analyzer, security):
    st.title("Data Cleaning")

    # Check if data is loaded
    if 'data' in session_state and session_state.data is not None:
        # Create a backup of the original data if not already present
        if 'original_data' not in session_state:
            session_state.original_data = session_state.data.copy()

        # Select cleaning operation
        cleaning_option = st.selectbox(
            "Select cleaning operation",
            ["Handle Missing Values", "Remove Duplicates", "Filter Data"]
        )

        # Handle Missing Values
        if cleaning_option == "Handle Missing Values":
            temp_data = session_state.data.copy()  # Temporary DataFrame for previewing changes
            st.write("Data with Missing Values:")
            st.write(temp_data[temp_data.isnull().any(axis=1)].head())  # Display rows with missing values

            # Handle missing values column by column
            for column in temp_data.columns:
                if temp_data[column].isnull().any():
                    st.write(f"Column with missing values: {column}")
                    method = st.radio(
                        f"Choose how to handle missing values in {column}:",
                        ["Drop", "Fill with mean", "Fill with median", "Fill with mode"],
                        key=column  # Ensure unique keys for each column
                    )

                    # Apply selected method to the temporary DataFrame
                    if method == "Drop":
                        temp_data = temp_data.dropna(subset=[column])
                    elif method == "Fill with mean" and pd.api.types.is_numeric_dtype(temp_data[column]):
                        temp_data[column].fillna(temp_data[column].mean(), inplace=True)
                    elif method == "Fill with median" and pd.api.types.is_numeric_dtype(temp_data[column]):
                        temp_data[column].fillna(temp_data[column].median(), inplace=True)
                    elif method == "Fill with mode":
                        temp_data[column].fillna(temp_data[column].mode()[0], inplace=True)

            # Display a preview of the cleaned data
            st.write("Preview of Cleaned Data:")
            st.write(temp_data.head())

            # Provide confirmation buttons
            if st.button("Confirm Changes"):
                session_state.data = temp_data
                st.success("Changes applied successfully!")
            if st.button("Undo Changes"):
                session_state.data = session_state.original_data.copy()
                st.success("Changes reverted to original data!")

        # Remove Duplicates
        elif cleaning_option == "Remove Duplicates":
            st.write(f"Current data shape: {session_state.data.shape}")
            if st.button("Remove Duplicate Rows"):
                session_state.data = session_state.data.drop_duplicates()
                st.success("Duplicate rows removed!")
                st.write(f"New data shape: {session_state.data.shape}")

        # Filter Data
        elif cleaning_option == "Filter Data":
            column = st.selectbox("Select column to filter", session_state.data.columns)
            if pd.api.types.is_numeric_dtype(session_state.data[column]):
                min_val, max_val = st.slider(
                    "Filter range",
                    float(session_state.data[column].min()),
                    float(session_state.data[column].max()),
                    (float(session_state.data[column].min()), float(session_state.data[column].max()))
                )
                temp_data = session_state.data[
                    (session_state.data[column] >= min_val) & (session_state.data[column] <= max_val)
                ]
                st.write("Filtered Data Preview:")
                st.write(temp_data.head())
                if st.button("Apply Filter"):
                    session_state.data = temp_data
                    st.success("Filter applied successfully!")
            else:
                st.warning("Selected column is not numeric. Filtering is only available for numeric columns.")

    else:
        # Warning if no data is loaded
        st.warning("Please load data first.")





# Load SpaCy model
nlp = spacy.load("en_core_web_sm")

def render_data_preprocessing(session_state, analyzer, security):
    st.title("Data Preprocessing")

    # Check if data is loaded
    if 'data' in session_state and session_state.data is not None:
        # Create a backup of the original data if not already present
        if 'preprocessing_backup' not in session_state:
            session_state.preprocessing_backup = session_state.data.copy()

        # Step 0: Allow the user to select the text column
        st.subheader("Select Text Column for Preprocessing")
        text_columns = session_state.data.select_dtypes(include=['object', 'string']).columns
        if len(text_columns) == 0:
            st.error("No text columns found in the dataset. Please upload a dataset with text data.")
            return

        # Dropdown to select text column
        selected_column = st.selectbox("Choose the column containing text data:", text_columns)
        st.write(f"Preview of selected column: {selected_column}")
        st.write(session_state.data[selected_column].head())

        # Step 1: Remove Irrelevant Elements
        if st.button("Remove Irrelevant Elements"):
            def clean_text(text):
                # Ensure input is a string
                if not isinstance(text, str):
                    text = str(text) if text is not None else ""
                # Remove URLs, mentions, hashtags, punctuation, and special characters
                text = re.sub(r"http\S+|www\S+|https\S+|\@\w+|\#|[^\w\s]", '', text)
                # Remove words with numbers and extra spaces
                text = re.sub(r'\w*\d\w*', ' ', text)
                text = re.sub('\s+', ' ', text).strip()
                return text.lower()

            # Apply the cleaning function
            cleaned_column = "cleaned_text"
            session_state.data[cleaned_column] = session_state.data[selected_column].apply(clean_text)

            st.success("Irrelevant elements removed successfully!")
            st.write(session_state.data[[selected_column, cleaned_column]].head())

        # Step 2: Tokenization
        if st.button("Tokenize Text"):
            def spacy_tokenizer(text):
                doc = nlp(text)
                return [token.text for token in doc if not token.is_space]

            cleaned_column = "cleaned_text"
            if cleaned_column in session_state.data.columns:
                session_state.data["tokens"] = session_state.data[cleaned_column].apply(spacy_tokenizer)

                st.success("Text tokenized successfully!")
                st.write(session_state.data[["tokens"]].head())
            else:
                st.error(f"'{cleaned_column}' does not exist. Please run 'Remove Irrelevant Elements' first.")

        # Step 3: Lemmatization
        if st.button("Perform Lemmatization"):
            def spacy_lemmatizer(text):
                doc = nlp(text)
                return [token.lemma_ for token in doc if not token.is_space]

            cleaned_column = "cleaned_text"
            if cleaned_column in session_state.data.columns:
                session_state.data["lemmatized_tokens"] = session_state.data[cleaned_column].apply(spacy_lemmatizer)

                st.success("Lemmatization completed successfully!")
                st.write(session_state.data[["tokens", "lemmatized_tokens"]].head())
            else:
                st.error(f"'{cleaned_column}' does not exist. Please run 'Remove Irrelevant Elements' first.")

        # Step 4: Keyword Filtering
        keywords = st.text_input("Enter keywords for filtering (comma-separated)", "nuclear, energy, renewable")
        if st.button("Apply Keyword Filtering"):
            lemmatized_column = "lemmatized_tokens"
            filtered_column = "filtered_data"

            if lemmatized_column in session_state.data.columns:
                keyword_list = [kw.strip().lower() for kw in keywords.split(",")]

                def filter_keywords(tokens):
                    # Check if any keyword exists in the list of tokens
                    return any(kw in tokens for kw in keyword_list)

                session_state.data[filtered_column] = session_state.data[lemmatized_column].apply(filter_keywords)
                session_state.data = session_state.data[session_state.data[filtered_column]]

                st.success("Keyword filtering applied successfully!")
                st.write(session_state.data.head())
            else:
                st.error(f"'{lemmatized_column}' does not exist. Please run 'Perform Lemmatization' first.")

        # Save or Undo Changes
        st.subheader("Actions")
        if st.button("Save Preprocessed Data"):
            session_state.preprocessing_backup = session_state.data.copy()
            st.success("Preprocessed data saved successfully!")

        if st.button("Undo Preprocessing Changes"):
            if 'preprocessing_backup' in session_state:
                session_state.data = session_state.preprocessing_backup.copy()
                st.success("Changes reverted to the last saved state!")
                st.write(session_state.data.head())
            else:
                st.error("No backup found. Please save preprocessed data before undoing changes.")
    else:
        st.error("No data loaded. Please upload a dataset to begin preprocessing.")



def get_column_by_name(session_state, possible_names):
    """Find a column in the dataset based on a list of possible names (case-insensitive)."""
    for col in session_state.data.columns:
        if col.lower() in [name.lower() for name in possible_names]:
            return col
    return None


def analyze_sentiment(text):
    """Perform sentiment analysis using SpaCy."""
    doc = nlp(text)
    polarity = 0
    for sentence in doc.sents:
        # Placeholder logic for sentiment analysis: you can replace this with a specific sentiment scoring library
        polarity += sentence.sentiment
    return "positive" if polarity > 0 else "negative" if polarity < 0 else "neutral"


def render_analysis(session_state, analyzer, security):
    # Ensure data is loaded
    if 'data' not in session_state or session_state.data is None:
        st.warning("Please load data first")
        return

    st.title("Data Analysis Methods and Visualization")

    # Detect relevant columns case-insensitively
    text_column = get_column_by_name(session_state, ["text", "Text"])
    sentiment_column = get_column_by_name(session_state, ["sentiment", "Sentiment"])

    # Check for sentiment column; if absent, generate it
    if sentiment_column is None:
        if text_column is not None:
            with st.spinner("Performing sentiment analysis..."):
                session_state.data['sentiment'] = session_state.data[text_column].apply(analyze_sentiment)
                sentiment_column = "sentiment"
                st.success("Sentiment analysis completed. 'sentiment' column added to dataset.")
        else:
            st.error("No text column found for sentiment analysis.")
            return

    # Let the user select analysis options from a dropdown
    analysis_type = st.selectbox(
        "Select a Data Analysis Method",
        [
            "Select Method",
            "Sentiment Distribution Analysis",
            "Word Frequency Analysis",
            "Temporal Sentiment Trend Analysis",
            "Keyword-Based Filtering and Analysis",
            "Model Performance Metrics",

        ]
    )

    # Perform Sentiment Distribution Analysis
    if analysis_type == "Sentiment Distribution Analysis":
        with st.spinner("Generating sentiment distribution..."):
            sentiment_counts = session_state.data[sentiment_column].value_counts()
            st.bar_chart(sentiment_counts)
        st.success("Sentiment distribution visualized!")

    # Perform Word Frequency Analysis
    if analysis_type == "Word Frequency Analysis":
        if text_column is not None:
            with st.spinner("Calculating word frequencies..."):
                sentiment = st.radio("Select Sentiment", ["positive", "negative", "neutral"])
                words = session_state.data.loc[session_state.data[sentiment_column] == sentiment, text_column]
                word_freq = pd.Series(' '.join(words).split()).value_counts()[:20]
                st.bar_chart(word_freq)
            st.success("Word frequency analysis completed!")
        else:
            st.warning("No text column found in the dataset.")

    # Generate Word Clouds
    if analysis_type == "Word Cloud Generation":
        if text_column is not None:
            with st.spinner("Generating word cloud..."):
                sentiment = st.radio("Select Sentiment", ["positive", "negative", "neutral"])
                words = session_state.data.loc[session_state.data[sentiment_column] == sentiment, text_column]
                wordcloud = WordCloud(width=800, height=400, background_color='white').generate(' '.join(words))
                st.image(wordcloud.to_array())
            st.success("Word cloud generated!")
        else:
            st.warning("No text column found in the dataset.")

    # Perform Temporal Sentiment Trend Analysis
    if analysis_type == "Temporal Sentiment Trend Analysis":
        if 'date' in session_state.data.columns:
            with st.spinner("Analyzing sentiment trends over time..."):
                session_state.data['date'] = pd.to_datetime(session_state.data['date'])
                trend_data = session_state.data.groupby([session_state.data['date'].dt.date, sentiment_column]).size().unstack()
                st.line_chart(trend_data)
            st.success("Sentiment trends analyzed!")
        else:
            st.warning("No 'date' column found in the dataset.")

    # Perform Keyword-Based Filtering and Analysis
    if analysis_type == "Keyword-Based Filtering and Analysis":
        if text_column is not None:
            with st.spinner("Filtering text by keywords..."):
                keywords = st.text_input("Enter keywords separated by commas (e.g., nuclear, renewable)").split(',')
                filtered_data = session_state.data[session_state.data[text_column].str.contains('|'.join(keywords), case=False, na=False)]
                filtered_sentiment_counts = filtered_data[sentiment_column].value_counts()
                st.bar_chart(filtered_sentiment_counts)
            st.success("Keyword-based filtering completed!")
        else:
            st.warning("No text column found in the dataset.")

    # Display Model Performance Metrics
    if analysis_type == "Model Performance Metrics":
        with st.spinner("Calculating model performance metrics..."):
            if 'true_labels' in session_state.data.columns and 'predicted_labels' in session_state.data.columns:
                from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score

                true_labels = session_state.data['true_labels']
                predicted_labels = session_state.data['predicted_labels']

                metrics = {
                    "Accuracy": accuracy_score(true_labels, predicted_labels),
                    "Precision": precision_score(true_labels, predicted_labels, average='weighted'),
                    "Recall": recall_score(true_labels, predicted_labels, average='weighted'),
                    "F1 Score": f1_score(true_labels, predicted_labels, average='weighted')
                }

                session_state['model_metrics'] = metrics
                st.write("Model Performance Metrics:")
                st.json(metrics)
            else:
                st.warning("The dataset must include `true_labels` and `predicted_labels` columns for model evaluation.")
        st.success("Model performance metrics displayed!")

    # Perform Clustering and Thematic Grouping
    if analysis_type == "Clustering and Thematic Grouping":
        if text_column is not None:
            with st.spinner("Performing clustering and thematic grouping..."):
                vectorizer = TfidfVectorizer(max_features=500)
                tfidf_matrix = vectorizer.fit_transform(session_state.data[text_column])
                kmeans = KMeans(n_clusters=3, random_state=42)
                clusters = kmeans.fit_predict(tfidf_matrix)
                session_state.data['cluster'] = clusters
                st.bar_chart(session_state.data['cluster'].value_counts())
            st.success("Clustering and thematic grouping completed!")
        else:
            st.warning("No text column found in the dataset.")

    # Preview analyzed data
    st.write("Preview of analyzed data:")
    st.dataframe(session_state.data.head())



def render_reports(session_state, analyzer, security):
    if 'data' not in session_state or session_state.data is None:
        st.warning("Please load data first")
        return

    # Normalize column names to lowercase for case-insensitive access
    session_state.data.columns = map(str.lower, session_state.data.columns)

    st.title("Generate Reports")

    # Dropdown for selecting the report to generate
    report_type = st.selectbox(
        "Select a Report to Generate",
        [
            "Select Report",
            "Sentiment Analysis Summary Report",
            "Top Keywords and Themes Report",
            "Temporal Insights Report",
            "Filtered Sentiment Report",
            "Model Evaluation Report"
        ]
    )

    # Generate and display the selected report
    if st.button("Generate Report"):
        with st.spinner(f"Generating {report_type}..."):
            report_data = {}
            report_content = None

            if report_type == "Sentiment Analysis Summary Report":
                sentiment_counts = session_state.data['sentiment'].value_counts().to_dict()
                st.bar_chart(session_state.data['sentiment'].value_counts())
                report_data = {
                    'timestamp': datetime.now().isoformat(),
                    'sentiment_distribution': sentiment_counts
                }
                report_content = f"Sentiment Distribution:\n{sentiment_counts}"

            elif report_type == "Top Keywords and Themes Report":
                words = session_state.data['text'].str.split().explode()
                top_keywords = words.value_counts().head(20).to_dict()
                import pandas as pd
                st.bar_chart(pd.Series(top_keywords))
                report_data = {
                    'timestamp': datetime.now().isoformat(),
                    'top_keywords': top_keywords
                }
                report_content = f"Top Keywords:\n{top_keywords}"

            elif report_type == "Temporal Insights Report":
                if 'date' in session_state.data.columns:
                    session_state.data['date'] = pd.to_datetime(session_state.data['date'])
                    temporal_data = session_state.data.groupby([session_state.data['date'].dt.date, 'sentiment']).size().unstack().fillna(0).to_dict()
                    st.line_chart(session_state.data.groupby([session_state.data['date'].dt.date, 'sentiment']).size().unstack())
                    report_data = {
                        'timestamp': datetime.now().isoformat(),
                        'temporal_trends': temporal_data
                    }
                    report_content = f"Temporal Trends:\n{temporal_data}"
                else:
                    st.warning("Date column is required for Temporal Insights Report.")
                    return

            elif report_type == "Filtered Sentiment Report":
                keywords = st.text_input("Enter keywords for filtering (e.g., nuclear, renewable)").split(',')
                if keywords:
                    filtered_data = session_state.data[session_state.data['text'].str.contains('|'.join(keywords), case=False, na=False)]
                    filtered_sentiment_counts = filtered_data['sentiment'].value_counts().to_dict()
                    import pandas as pd
                    st.bar_chart(pd.Series(filtered_sentiment_counts))
                    report_data = {
                        'timestamp': datetime.now().isoformat(),
                        'filtered_sentiment_distribution': filtered_sentiment_counts
                    }
                    report_content = f"Filtered Sentiment Distribution:\n{filtered_sentiment_counts}"
                else:
                    st.warning("Enter valid keywords for filtering.")
                    return

            elif report_type == "Model Evaluation Report":
                if 'model_metrics' in session_state:
                    model_metrics = session_state['model_metrics']  # Access dynamically stored model metrics
                else:
                    st.warning("Model evaluation metrics are missing. Please ensure they are calculated in previous steps.")
                    return

                st.write(model_metrics)
                report_data = {
                    'timestamp': datetime.now().isoformat(),
                    'model_metrics': model_metrics
                }
                report_content = f"Model Evaluation Metrics:\n{model_metrics}"

            else:
                st.warning("Please select a valid report type.")
                return

            st.success(f"{report_type} generated successfully!")

            # Option to save the report
            file_format = st.radio("Select file format to save the report:", ["PDF", "Word", "Excel"])
            if st.button("Save Report"):
                if file_format == "PDF":
                    from fpdf import FPDF
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    pdf.multi_cell(0, 10, report_content)
                    filename = f"{report_type.replace(' ', '_').lower()}.pdf"
                    pdf.output(filename)
                elif file_format == "Word":
                    import docx
                    doc = docx.Document()
                    doc.add_heading(report_type, level=1)
                    doc.add_paragraph(report_content)
                    filename = f"{report_type.replace(' ', '_').lower()}.docx"
                    doc.save(filename)
                elif file_format == "Excel":
                    import pandas as pd
                    df = pd.DataFrame.from_dict(report_data, orient='index')
                    filename = f"{report_type.replace(' ', '_').lower()}.xlsx"
                    df.to_excel(filename, index=False)
                st.success(f"Report saved as {filename}.")


def render_help(session_state, analyzer, security):
    st.title("Help & Documentation")

    st.markdown("""
    ### Analysis Methods
    - **Sentiment Analysis**: Uses VADER and TextBlob for sentiment scoring
    - **Text Statistics**: Basic text metrics and patterns

    ### Data Format Requirements
    - CSV files with a 'text' column
    - Plain text files (one entry per line)
    - Direct text input

    ### Security Features
    - File validation and size limits
    - Input sanitization
    - Basic authentication
    """)

if __name__ == "__main__":
    main()
